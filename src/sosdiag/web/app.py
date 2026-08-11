from __future__ import annotations

import shutil
import uuid
from html import escape
from pathlib import Path
from typing import Annotated

from docx import Document
from fastapi import FastAPI, File, Form, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from jinja2 import Environment, FileSystemLoader, select_autoescape

APP_ROOT = Path(__file__).resolve().parents[3]
DATA_DIR = Path("/app/data")
UPLOAD_DIR = DATA_DIR / "uploads"
OUTPUT_DIR = DATA_DIR / "output"
WEB_TEMPLATE_DIR = APP_ROOT / "templates" / "web"

for directory in (UPLOAD_DIR, OUTPUT_DIR):
    directory.mkdir(parents=True, exist_ok=True)

env = Environment(
    loader=FileSystemLoader(WEB_TEMPLATE_DIR),
    autoescape=select_autoescape(["html", "xml"]),
)

app = FastAPI(title="sosreport Structure Diagnostic")


@app.get("/", response_class=HTMLResponse)
def index() -> HTMLResponse:
    template = env.get_template("index.html.j2")
    return HTMLResponse(template.render())


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/reports", response_class=HTMLResponse)
async def create_report(
    sosreports: Annotated[list[UploadFile], File(...)],
    customer_name: Annotated[str, Form(...)],
    execution_period: Annotated[str, Form(...)],
    location: Annotated[str, Form(...)],
    customer_contact: Annotated[str, Form("")],
    sales_name: Annotated[str, Form("")],
    sales_role: Annotated[str, Form("")],
    sales_phone: Annotated[str, Form("")],
    sales_email: Annotated[str, Form("")],
    engineer_name: Annotated[str, Form("")],
    engineer_role: Annotated[str, Form("")],
    engineer_phone: Annotated[str, Form("")],
    engineer_email: Annotated[str, Form("")],
    output_format: Annotated[str, Form("html")],
) -> HTMLResponse:
    job_id = uuid.uuid4().hex[:12]
    job_upload_dir = UPLOAD_DIR / job_id
    job_output_dir = OUTPUT_DIR / job_id
    job_upload_dir.mkdir(parents=True, exist_ok=True)
    job_output_dir.mkdir(parents=True, exist_ok=True)

    saved_files: list[str] = []
    for upload in sosreports:
        safe_name = Path(upload.filename or "sosreport").name
        target = job_upload_dir / safe_name
        with target.open("wb") as fh:
            shutil.copyfileobj(upload.file, fh)
        saved_files.append(safe_name)

    metadata = {
        "customer_name": customer_name,
        "execution_period": execution_period,
        "location": location,
        "customer_contact": customer_contact,
        "sales": {
            "name": sales_name,
            "role": sales_role,
            "phone": sales_phone,
            "email": sales_email,
        },
        "engineers": [
            {
                "name": engineer_name,
                "role": engineer_role,
                "phone": engineer_phone,
                "email": engineer_email,
            }
        ],
    }

    # Production integration point:
    # 1. validate and safely extract archives
    # 2. parse each sosreport
    # 3. evaluate YAML rules
    # 4. aggregate a DiagnosticReport
    # 5. pass that same model to HTML and DOCX renderers
    generated: list[dict[str, str]] = []

    if output_format in {"html", "both"}:
        html_path = job_output_dir / "structure-diagnostic.html"
        _write_placeholder_html(html_path, job_id, metadata, saved_files)
        generated.append(
            {
                "name": html_path.name,
                "url": f"/reports/{job_id}/{html_path.name}",
            }
        )

    if output_format in {"docx", "both"}:
        docx_path = job_output_dir / "structure-diagnostic.docx"
        _write_placeholder_docx(docx_path, job_id, metadata, saved_files)
        generated.append(
            {
                "name": docx_path.name,
                "url": f"/reports/{job_id}/{docx_path.name}",
            }
        )

    template = env.get_template("result.html.j2")
    return HTMLResponse(
        template.render(
            job_id=job_id,
            customer_name=customer_name,
            uploaded=saved_files,
            generated=generated,
        )
    )


@app.get("/reports/{job_id}/{filename}")
def download_report(job_id: str, filename: str) -> FileResponse:
    safe_job = Path(job_id).name
    safe_name = Path(filename).name
    target = OUTPUT_DIR / safe_job / safe_name
    return FileResponse(target, filename=safe_name)


def _write_placeholder_html(
    path: Path, job_id: str, metadata: dict, uploaded: list[str]
) -> None:
    uploaded_rows = "".join(f"<li>{escape(name)}</li>" for name in uploaded)
    sales = metadata["sales"]
    engineer = metadata["engineers"][0]
    path.write_text(
        f"""<!doctype html>
<html lang=\"ko\">
<head>
<meta charset=\"utf-8\">
<title>{escape(metadata['customer_name'])} Health Check</title>
<style>
body{{font-family:Arial,sans-serif;max-width:980px;margin:40px auto;padding:0 24px;color:#222}}
h1{{border-bottom:3px solid #0f4c81;padding-bottom:12px}}table{{width:100%;border-collapse:collapse;margin:14px 0 26px}}th,td{{border:1px solid #999;padding:8px}}th{{background:#e8eef5}}.notice{{padding:14px;border-left:5px solid #0f4c81;background:#f5f8fb}}
</style>
</head>
<body>
<h1>{escape(metadata['customer_name'])} Health Check 보고서</h1>
<h2>1. 구조진단 개요</h2>
<h3>1.1 수행 정보</h3>
<table><tr><th>고객사</th><td>{escape(metadata['customer_name'])}</td><th>기간</th><td>{escape(metadata['execution_period'])}</td></tr><tr><th>장소</th><td>{escape(metadata['location'])}</td><th>고객 담당자</th><td>{escape(metadata['customer_contact'])}</td></tr></table>
<h3>1.2 rockPLACE 영업 대표</h3>
<table><tr><th>Name</th><th>Role</th><th>Phone</th><th>E-Mail</th></tr><tr><td>{escape(sales['name'])}</td><td>{escape(sales['role'])}</td><td>{escape(sales['phone'])}</td><td>{escape(sales['email'])}</td></tr></table>
<h3>1.3 rockPLACE 기술 담당</h3>
<table><tr><th>Name</th><th>Role</th><th>Phone</th><th>E-Mail</th></tr><tr><td>{escape(engineer['name'])}</td><td>{escape(engineer['role'])}</td><td>{escape(engineer['phone'])}</td><td>{escape(engineer['email'])}</td></tr></table>
<h3>1.4 업로드된 대상 시스템 자료</h3><ul>{uploaded_rows}</ul>
<div class=\"notice\"><strong>Job ID:</strong> {job_id}<br>현재 파일은 웹 업로드/출력 흐름 검증용 보고서입니다. 전체 sosreport 진단 엔진이 연결되면 현재 진단 항목의 A/B/C 결과와 상세 Report가 이 문서에 생성됩니다.</div>
</body></html>""",
        encoding="utf-8",
    )


def _write_placeholder_docx(
    path: Path, job_id: str, metadata: dict, uploaded: list[str]
) -> None:
    document = Document()
    document.add_heading(f"{metadata['customer_name']} Health Check 보고서", level=0)
    document.add_paragraph("Red Hat Enterprise Linux")

    document.add_heading("1. 구조진단 개요", level=1)
    document.add_heading("1.1 수행 정보", level=2)
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    values = [
        ("고객사", metadata["customer_name"], "기간", metadata["execution_period"]),
        ("장소", metadata["location"], "고객 담당자", metadata["customer_contact"]),
    ]
    for row, values_row in zip(table.rows, values):
        for cell, value in zip(row.cells, values_row):
            cell.text = str(value)

    document.add_heading("1.2 rockPLACE 영업 대표", level=2)
    _add_person_table(document, metadata["sales"])

    document.add_heading("1.3 rockPLACE 기술 담당", level=2)
    _add_person_table(document, metadata["engineers"][0])

    document.add_heading("1.4 업로드된 대상 시스템 자료", level=2)
    for name in uploaded:
        document.add_paragraph(name, style="List Bullet")

    document.add_heading("2. 상세 Report", level=1)
    document.add_paragraph(
        "현재 파일은 웹 업로드/출력 흐름 검증용 보고서입니다. "
        "전체 sosreport 진단 엔진이 연결되면 현재 진단 항목의 A/B/C 결과, "
        "검토 의견, 점검 내역 및 조치 사항이 이 영역에 생성됩니다."
    )
    document.add_paragraph(f"Job ID: {job_id}")
    document.save(path)


def _add_person_table(document: Document, person: dict) -> None:
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["Name", "Role", "Phone", "E-Mail"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for cell, key in zip(table.rows[1].cells, ["name", "role", "phone", "email"]):
        cell.text = str(person.get(key, ""))


def main() -> None:
    import uvicorn

    uvicorn.run("sosdiag.web.app:app", host="0.0.0.0", port=8000)


if __name__ == "__main__":
    main()
